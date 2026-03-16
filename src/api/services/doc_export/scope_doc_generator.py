"""Generates annotated scope letter .docx files from session match results."""

import io
import logging
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from scope_classification import ConnectionFactory

from ._dataclasses import HighlightRegion
from .run_splitter import RunSplitter


log = logging.getLogger(__name__)

# ── Production colors (by match type) ────────────────────────────────
ALIGNED_COLOR = WD_COLOR_INDEX.BRIGHT_GREEN
PARTIAL_COLOR = WD_COLOR_INDEX.YELLOW

# FIXME: dev-only — low confidence mapping indicator, remove before release
LOW_CONFIDENCE_COLOR    = WD_COLOR_INDEX.TURQUOISE
LOW_CONFIDENCE_THRESHOLD = 0.60

# ── Verification colors (by match method) ────────────────────────────
VERIFY_EXACT_COLOR  = WD_COLOR_INDEX.BRIGHT_GREEN
VERIFY_FUZZY_COLOR  = WD_COLOR_INDEX.YELLOW
VERIFY_MANUAL_COLOR = WD_COLOR_INDEX.PINK


class ScopeDocGenerator:
    """Generates highlighted scope letter documents from analysis results."""

    def __init__(self, db: ConnectionFactory, template_path: str) -> None:
        """Initialize with DB connection and path to template .docx."""

        self._db            = db
        self._schema        = db.schema
        self._template_path = template_path
        self._splitter      = RunSplitter()


    def generate_session_doc(self, session_id: int) -> tuple[bytes, str]:
        """Generate highlighted scope letter for a completed session. Returns (docx_bytes, filename)."""

        session  = self._load_session(session_id)
        matches  = self._load_session_matches(session_id)
        mappings = self._load_template_mappings()

        # Build MfcExclusionId -> match type lookup
        match_lookup = {}
        for m in matches:
            mfc_id     = m['MfcExclusionId']
            match_type = m['MatchType']
            if mfc_id and match_type in ('Aligned', 'Partial'):
                # Conservative: Partial wins if same MFC item matched multiple ways
                existing = match_lookup.get(mfc_id)
                if existing == 'Partial' or match_type == 'Partial':
                    match_lookup[mfc_id] = 'Partial'
                else:
                    match_lookup[mfc_id] = 'Aligned'

        # Build mapping lookup: MfcExclusionId -> mapping row
        mapping_lookup = {m['MfcExclusionId']: m for m in mappings}

        # Build highlight regions grouped by paragraph index
        regions_by_para = defaultdict(list)

        for mfc_id, match_type in match_lookup.items():
            mapping = mapping_lookup.get(mfc_id)
            if not mapping:
                log.warning(f"  MFC exclusion {mfc_id}: matched as {match_type} but no template mapping exists")
                continue

            # Determine color
            if match_type == 'Partial':
                color = PARTIAL_COLOR
            elif float(mapping['MatchConfidence']) < LOW_CONFIDENCE_THRESHOLD:
                # FIXME: dev-only blue highlight for low-confidence mappings
                color = LOW_CONFIDENCE_COLOR
            else:
                color = ALIGNED_COLOR

            region = HighlightRegion(
                start_char       = mapping['StartChar'],
                end_char         = mapping['EndChar'],
                color            = color,
                mfc_exclusion_id = mfc_id,
            )

            regions_by_para[mapping['ParaIndex']].append(region)

        # Load template and apply highlights
        doc       = Document(self._template_path)
        total_hl  = 0

        for para_idx, regions in regions_by_para.items():
            if para_idx >= len(doc.paragraphs):
                log.warning(f"  Para index {para_idx} out of range (doc has {len(doc.paragraphs)} paragraphs)")
                continue

            paragraph = doc.paragraphs[para_idx]
            applied   = self._splitter.apply_highlights(paragraph, regions)
            total_hl += applied

        log.info(f"Session {session_id}: highlighted {total_hl} runs across {len(regions_by_para)} paragraphs")

        # Build filename
        filename = self._build_filename(session)

        # Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return buf.getvalue(), filename


    def generate_verification_doc(self) -> bytes:
        """
        DEV: Verify template mapping positions.

        Highlights all mapped exclusions colored by match method:
        Green = Exact, Yellow = Fuzzy, Red = Manual, Blue = low confidence.
        This method is used to verify template match locations/output.
        """

        mappings = self._load_template_mappings()

        log.info(f"Verification: loaded {len(mappings)} template mappings")

        regions_by_para = defaultdict(list)

        for mapping in mappings:
            method     = mapping['MatchMethod']
            confidence = float(mapping['MatchConfidence'])

            # FIXME: dev-only — blue for low confidence regardless of method
            if confidence < LOW_CONFIDENCE_THRESHOLD:
                color = LOW_CONFIDENCE_COLOR
            elif method == 'Exact':
                color = VERIFY_EXACT_COLOR
            elif method == 'Fuzzy':
                color = VERIFY_FUZZY_COLOR
            else:
                color = VERIFY_MANUAL_COLOR

            region = HighlightRegion(
                start_char       = mapping['StartChar'],
                end_char         = mapping['EndChar'],
                color            = color,
                mfc_exclusion_id = mapping['MfcExclusionId'],
            )

            regions_by_para[mapping['ParaIndex']].append(region)

        log.info(f"Verification: {len(regions_by_para)} paragraphs to highlight")

        doc      = Document(self._template_path)
        total_hl = 0

        for para_idx, regions in regions_by_para.items():
            if para_idx >= len(doc.paragraphs):
                log.warning(f"  Verify: para index {para_idx} out of range (doc has {len(doc.paragraphs)})")
                continue

            paragraph = doc.paragraphs[para_idx]
            para_text = paragraph.text
            log.debug(f"  Para {para_idx}: {len(regions)} regions, text length {len(para_text)}")

            applied   = self._splitter.apply_highlights(paragraph, regions)
            total_hl += applied

        log.info(f"Verification doc: highlighted {total_hl} chars across {len(regions_by_para)} paragraphs")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return buf.getvalue()


    def generate_editor_export(
            self,
            session_id: int,
            view_mode:  str = 'full' ) -> tuple[bytes, str]:
        """Generate a clean (no highlights) .docx respecting editor state."""

        session        = self._load_session(session_id)
        mappings       = self._load_template_mappings()
        section_ranges = self._load_section_ranges()
        editor_state   = self._load_editor_state(session_id)

        # Build lookup sets from editor state
        removed_paras   = set(editor_state['removed_paragraphs'])
        removed_regions = {
            (r['para_index'], r['mfc_exclusion_id'])
            for r in editor_state['removed_regions']
        }
        text_edits = {
            e['para_index']: e['edited_text']
            for e in editor_state['text_edits']
        }

        # Build mapping lookup: para_index -> list of mappings
        mappings_by_para: dict[int, list[dict]] = defaultdict(list)
        for mp in mappings:
            mappings_by_para[mp['ParaIndex']].append(mp)

        # Determine which paragraphs are visible in the current view
        def para_in_view(idx: int) -> bool:
            if view_mode == 'full':
                return True

            for section_name, (range_start, range_end) in section_ranges.items():
                if section_name == view_mode and range_start <= idx <= range_end:
                    return True

            return False

        # Load template and process
        doc          = Document(self._template_path)
        paras_to_del = []

        for idx, para in enumerate(doc.paragraphs):
            # Filter by view mode
            if not para_in_view(idx):
                paras_to_del.append(para)
                continue

            # Remove deleted paragraphs
            if idx in removed_paras:
                paras_to_del.append(para)
                continue

            # Apply text edits (user typed replacement)
            if idx in text_edits:
                self._replace_para_text(para, text_edits[idx])
                continue

            # Remove excluded regions (strip char ranges from text)
            para_mappings = mappings_by_para.get(idx, [])
            removed_spans = []
            for mp in para_mappings:
                key = (idx, mp['MfcExclusionId'])
                if key in removed_regions:
                    removed_spans.append((mp['StartChar'], mp['EndChar']))

            if removed_spans:
                self._strip_char_ranges(para, removed_spans)

        # Delete marked paragraphs (reverse order to preserve indices)
        for para in reversed(paras_to_del):
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)

        # Build filename
        filename = self._build_filename(session).replace('_Scope_Analysis', '_Scope_Letter_Edited')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return buf.getvalue(), filename


    @staticmethod
    def _replace_para_text(para, new_text: str) -> None:
        """Replace all runs in a paragraph with a single run containing new text."""

        # Preserve formatting from first run
        first_run  = para.runs[0] if para.runs else None
        font_name  = first_run.font.name if first_run else None
        font_size  = first_run.font.size if first_run else None
        bold       = first_run.bold if first_run else None

        # Clear all existing runs
        for run in para.runs:
            run._element.getparent().remove(run._element)

        # Add new run with text
        new_run = para.add_run(new_text)
        if font_name: new_run.font.name = font_name
        if font_size: new_run.font.size = font_size
        if bold:      new_run.bold      = bold


    @staticmethod
    def _strip_char_ranges(para, spans: list[tuple[int, int]]) -> None:
        """Remove character ranges from a paragraph, preserving remaining text and formatting."""

        # Sort spans and merge overlapping
        spans.sort(key=lambda s: s[0])
        merged = []
        for start, end in spans:
            if merged and start <= merged[-1][1]:
                merged[-1] = (merged[-1][0], max(merged[-1][1], end))
            else:
                merged.append((start, end))

        # Build set of character positions to remove
        remove_chars = set()
        for start, end in merged:
            for i in range(start, end):
                remove_chars.add(i)

        if not remove_chars:
            return

        # Walk runs, rebuild each run's text with removed chars stripped
        cum_pos = 0
        for run in para.runs:
            run_text = run.text or ''
            if not run_text:
                continue

            new_chars = []
            for offset, char in enumerate(run_text):
                if (cum_pos + offset) not in remove_chars:
                    new_chars.append(char)

            run.text = ''.join(new_chars)
            cum_pos += len(run_text)


    def build_editor_data(self, session_id: int) -> dict:
        """Build structured JSON for the browser scope letter editor."""

        session        = self._load_session(session_id)
        matches        = self._load_session_matches_full(session_id)
        mappings       = self._load_template_mappings()
        section_ranges = self._load_section_ranges()
        editor_state   = self._load_editor_state(session_id)

        # Build MfcExclusionId -> best match row
        match_lookup: dict[int, dict] = {}
        for m in matches:
            mfc_id = m['MfcExclusionId']
            if not mfc_id:
                continue
            existing = match_lookup.get(mfc_id)
            if not existing or m['MatchType'] == 'Partial':
                match_lookup[mfc_id] = m

        # Build mapping lookup keyed by para index
        mappings_by_para: dict[int, list[dict]] = defaultdict(list)
        for mp in mappings:
            mappings_by_para[mp['ParaIndex']].append(mp)

        # Load template and build paragraph data with run formatting
        doc        = Document(self._template_path)
        paragraphs = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text

            # Paragraph-level formatting
            pf         = para.paragraph_format
            indent_emu = pf.left_indent
            indent_in  = round(indent_emu / 914400, 2) if indent_emu else None

            # Hanging indent: first line is outdented from left_indent
            # In Word XML: w:ind left="2160" hanging="2160" means:
            #   continuation lines at 1.5", first line at 0"
            hanging_emu  = None
            first_ln_emu = None
            pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                ind = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind')
                if ind is not None:
                    hang_val = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging')
                    first_val = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine')
                    left_val = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')

                    if hang_val:
                        hanging_emu = int(hang_val) * 635  # twips to EMU
                    if first_val:
                        first_ln_emu = int(first_val) * 635
                    if left_val:
                        indent_emu = int(left_val) * 635
                        indent_in  = round(indent_emu / 914400, 2)

            hanging_in   = round(hanging_emu / 914400, 2)  if hanging_emu  else None
            first_line_in = round(first_ln_emu / 914400, 2) if first_ln_emu else None

            # Alignment: 0=LEFT, 1=CENTER, 2=RIGHT, 3=JUSTIFY
            alignment = None
            if para.alignment is not None:
                align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
                alignment = align_map.get(int(para.alignment), None)

            # Spacing before/after in points
            space_before_pt = round(pf.space_before / 12700, 1) if pf.space_before else None
            space_after_pt  = round(pf.space_after / 12700, 1)  if pf.space_after else None

            # Line spacing
            line_spacing = None
            if pf.line_spacing is not None:
                if pf.line_spacing_rule is not None and int(pf.line_spacing_rule) == 0:
                    # Single / 1.5 / Double — stored as a multiple
                    line_spacing = float(pf.line_spacing)
                elif pf.line_spacing:
                    # Exact or At Least — stored in EMU
                    line_spacing_pt = round(pf.line_spacing / 12700, 1)
                    line_spacing = line_spacing_pt

            # Tab stops
            tab_stops = []
            if pf.tab_stops:
                for ts in pf.tab_stops:
                    tab_stops.append({
                        'position_in': round(ts.position / 914400, 2),
                        'alignment':   str(ts.alignment).split('(')[0].strip() if ts.alignment else 'LEFT',
                    })

            # Build region lookup: char position -> region data
            region_map  = {}   # start_char -> region dict
            region_ends = {}   # char_pos   -> True if inside a region
            para_regions = []

            for mp in mappings_by_para.get(idx, []):
                mfc_id = mp['MfcExclusionId']
                match  = match_lookup.get(mfc_id)

                region = {
                    'mfc_id':       mfc_id,
                    'start':        mp['StartChar'],
                    'end':          mp['EndChar'],
                    'snippet':      mp['TemplateSnippet'],
                    'match_type':   match['MatchType'] if match else None,
                    'confidence':   float(match['Confidence']) if match and match['Confidence'] else None,
                    'risk_level':   match['RiskLevel'] if match else None,
                    'risk_notes':   match['RiskNotes'] if match else None,
                    'ai_reasoning': match['AiReasoning'] if match else None,
                    'erector_text': match['ErectorText'] if match else None,
                }

                region_map[mp['StartChar']] = region
                para_regions.append(region)

                for ci in range(mp['StartChar'], mp['EndChar']):
                    region_ends[ci] = region

            # Build segments by walking runs and splitting at region boundaries
            segments = []
            cum_pos  = 0

            for run in para.runs:
                run_text = run.text or ''
                if not run_text:
                    continue

                bold      = run.bold or False
                italic    = run.italic or False
                underline = run.underline or False
                size_pt   = round(run.font.size / 12700, 1) if run.font.size else None
                color_hex = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None

                # Split this run at region boundaries
                cur_chars    = []
                cur_region   = region_ends.get(cum_pos)

                for offset, char in enumerate(run_text):
                    char_pos   = cum_pos + offset
                    char_region = region_ends.get(char_pos)

                    if char_region is not cur_region and cur_chars:
                        segments.append(self._make_segment(
                            ''.join(cur_chars), bold, italic, underline, size_pt, color_hex, cur_region,
                        ))
                        cur_chars = []

                    cur_region = char_region
                    cur_chars.append(char)

                if cur_chars:
                    segments.append(self._make_segment(
                        ''.join(cur_chars), bold, italic, underline, size_pt, color_hex, cur_region,
                    ))

                cum_pos += len(run_text)

            # Determine which template section this paragraph belongs to
            template_section = None
            for section_name, (range_start, range_end) in section_ranges.items():
                if range_start <= idx <= range_end:
                    template_section = section_name
                    break

            paragraphs.append({
                'index':            idx,
                'text':             text,
                'indent':           indent_in,
                'hanging':          hanging_in,
                'first_line':       first_line_in,
                'alignment':        alignment,
                'space_before':     space_before_pt,
                'space_after':      space_after_pt,
                'line_spacing':     line_spacing,
                'tab_stops':        tab_stops if tab_stops else None,
                'segments':         segments,
                'regions':          para_regions,
                'template_section': template_section,
            })

        return {
            'session': {
                'id':       session.get('Id'),
                'erector':  session.get('ErectorNameRaw'),
                'job':      session.get('JobNumber'),
                'job_name': session.get('JobName'),
            },
            'paragraphs':   paragraphs,
            'editor_state': editor_state,
        }


    @staticmethod
    def _make_segment(
            text: str, bold: bool, italic: bool,
            underline: bool, size_pt: float | None,
            color: str | None, region: dict | None ) -> dict:
        """Build a single segment dict for the editor."""

        seg: dict = {'text': text}

        if bold:      seg['bold']      = True
        if italic:    seg['italic']    = True
        if underline: seg['underline'] = True
        if size_pt:   seg['size']      = size_pt
        if color:     seg['color']     = f'#{color}'

        if region:
            seg['region'] = {
                'mfc_id':     region['mfc_id'],
                'match_type': region['match_type'],
                'risk_level': region['risk_level'],
            }

        return seg


    # ── DB queries ───────────────────────────────────────────────────

    def _load_session(self, session_id: int) -> dict:
        """Load session metadata."""

        sql    = f"SELECT * FROM {self._schema}.AnalysisSessions WHERE Id = ?"
        cursor = self._db.execute(sql, (session_id,))
        row    = cursor.fetchone()

        if not row:
            raise ValueError(f"Session {session_id} not found")

        cols = [col[0] for col in cursor.description]

        return dict(zip(cols, row))


    def _load_session_matches(self, session_id: int) -> list[dict]:
        """Load all match results for a session."""

        sql = f"""
            SELECT MfcExclusionId, MatchType
            FROM {self._schema}.ExclusionMatches
            WHERE SessionId = ? AND MfcExclusionId IS NOT NULL
        """

        cursor = self._db.execute(sql, (session_id,))
        cols   = [col[0] for col in cursor.description]

        return [dict(zip(cols, row)) for row in cursor.fetchall()]


    def _load_session_matches_full(self, session_id: int) -> list[dict]:
        """Load match results with risk, reasoning, and erector text."""

        sql = f"""
            SELECT em.MfcExclusionId, em.MatchType, em.Confidence,
                   em.RiskLevel, em.RiskNotes, em.AiReasoning,
                   ee.RawText AS ErectorText
            FROM {self._schema}.ExclusionMatches em
            LEFT JOIN {self._schema}.ExtractedExclusions ee ON ee.Id = em.ExtractedExclusionId
            WHERE em.SessionId = ? AND em.MfcExclusionId IS NOT NULL
        """

        cursor = self._db.execute(sql, (session_id,))
        cols   = [col[0] for col in cursor.description]

        return [dict(zip(cols, row)) for row in cursor.fetchall()]


    def _load_template_mappings(self) -> list[dict]:
        """Load all template mappings."""

        sql = f"""
            SELECT MfcExclusionId, ParaIndex, StartChar, EndChar,
                   TemplateSnippet, MatchMethod, MatchConfidence
            FROM {self._schema}.TemplateMappings
            ORDER BY ParaIndex, StartChar
        """

        cursor = self._db.execute(sql)
        cols   = [col[0] for col in cursor.description]

        return [dict(zip(cols, row)) for row in cursor.fetchall()]


    def _load_editor_state(self, session_id: int) -> dict:
        """Load persisted editor state for a session."""

        schema = self._schema

        # Removed regions
        cursor = self._db.execute(
            f"SELECT MfcExclusionId, ParaIndex FROM {schema}.EditorRemovedRegions WHERE SessionId = ?",
            (session_id,),
        )
        removed_regions = [
            {'mfc_exclusion_id': row[0], 'para_index': row[1]}
            for row in cursor.fetchall()
        ]

        # Removed paragraphs
        cursor = self._db.execute(
            f"SELECT ParaIndex FROM {schema}.EditorRemovedParagraphs WHERE SessionId = ?",
            (session_id,),
        )
        removed_paragraphs = [row[0] for row in cursor.fetchall()]

        # Text edits
        cursor = self._db.execute(
            f"SELECT ParaIndex, EditedText FROM {schema}.EditorTextEdits WHERE SessionId = ?",
            (session_id,),
        )
        text_edits = [
            {'para_index': row[0], 'edited_text': row[1]}
            for row in cursor.fetchall()
        ]

        return {
            'removed_regions':    removed_regions,
            'removed_paragraphs': removed_paragraphs,
            'text_edits':         text_edits,
        }


    def _load_section_ranges(self) -> dict[str, tuple[int, int]]:
        """Load paragraph index ranges for each template section by ScopeType."""

        sql = f"""
            SELECT me.ScopeType,
                   MIN(tm.ParaIndex) AS MinPara,
                   MAX(tm.ParaIndex) AS MaxPara
            FROM {self._schema}.TemplateMappings tm
            JOIN {self._schema}.MfcExclusions me ON me.Id = tm.MfcExclusionId
            GROUP BY me.ScopeType
        """

        cursor = self._db.execute(sql)
        ranges = {}

        for row in cursor.fetchall():
            scope_type = row[0]
            min_para   = row[1]
            max_para   = row[2]

            if scope_type == 'Erect':
                # Include the section header paragraph (one above first mapped para)
                ranges['erector_exclusions'] = (min_para - 1, max_para)

        return ranges


    # ── Helpers ──────────────────────────────────────────────────────

    @staticmethod
    def _build_filename(session: dict) -> str:
        """Build output filename from session metadata."""

        job_number   = session.get('JobNumber') or 'NoJob'
        erector_name = session.get('ErectorNameRaw') or 'Unknown'

        # Sanitize for filesystem
        safe_job     = re.sub(r'[^\w\-]', '_', job_number.strip())
        safe_erector = re.sub(r'[^\w\-]', '_', erector_name.strip())

        return f"{safe_job}_{safe_erector}_Scope_Analysis.docx"
